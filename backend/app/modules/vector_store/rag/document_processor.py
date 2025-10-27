"""Document processor using LangChain for RAG."""

import io
import re
from pathlib import Path

import chardet
import docx
import PyPDF2
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.core.logger import logger


class DocumentProcessor:
    """Process documents (PDF, DOCX, TXT) using LangChain for RAG embeddings."""

    SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx", ".md"}
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB

    def __init__(self):
        # Initialize LangChain text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            is_separator_regex=False,
            separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""],
        )

    def is_supported(self, filename: str) -> bool:
        """Check if file type is supported."""
        extension = Path(filename).suffix.lower()
        return extension in self.SUPPORTED_EXTENSIONS

    def validate_file_size(self, file_size: int) -> tuple[bool, str | None]:
        """Validate file size is within limits."""
        if file_size > self.MAX_FILE_SIZE:
            max_mb = self.MAX_FILE_SIZE / (1024 * 1024)
            current_mb = file_size / (1024 * 1024)
            return False, f"File size {current_mb:.1f}MB exceeds limit of {max_mb}MB"
        return True, None

    def extract_text(self, file_content: bytes, filename: str) -> str:
        """Extract text content from file bytes with error recovery."""
        file_extension = Path(filename).suffix.lower()

        try:
            if file_extension == ".pdf":
                text = self._extract_from_pdf(file_content)
            elif file_extension == ".docx":
                text = self._extract_from_docx(file_content)
            elif file_extension in [".txt", ".md"]:
                text = self._extract_from_text(file_content)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")

            # Validate extracted text
            if not text or len(text.strip()) < 10:
                raise ValueError(f"Extracted text too short from {filename} (may be empty or corrupted)")

            return text

        except ValueError:
            raise
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {str(e)}")
            raise ValueError(f"Extraction failed: {str(e)}")

    def _extract_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file with error recovery."""
        try:
            # Try with strict=False to handle corrupted PDFs
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content), strict=False)
            pages_text = []

            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text and text.strip():
                        pages_text.append(text)
                except Exception as page_error:
                    logger.warning(f"Failed to extract page {page_num + 1}: {str(page_error)}")
                    continue

            full_text = "\n\n".join(pages_text)

            if not full_text.strip():
                raise ValueError("No text content extracted from PDF (empty or corrupted)")

            logger.info(f"Extracted text from {len(pages_text)} PDF pages")
            return full_text

        except Exception as e:
            logger.error(f"Error extracting PDF: {str(e)}")
            raise ValueError(f"PDF extraction failed: {str(e)}")

    def _extract_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file."""
        try:
            doc = docx.Document(io.BytesIO(file_content))
            paragraphs = []

            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)

            full_text = "\n\n".join(paragraphs)

            if not full_text.strip():
                raise ValueError("No text content extracted from DOCX")

            return full_text

        except Exception as e:
            logger.error(f"Error extracting DOCX: {str(e)}")
            raise ValueError(f"Failed to extract DOCX content: {str(e)}")

    def _extract_from_text(self, file_content: bytes) -> str:
        """Extract text from text-based files with encoding detection."""
        try:
            # Detect encoding
            detected = chardet.detect(file_content)
            encoding = detected.get("encoding", "utf-8")

            # Try detected encoding first
            try:
                text = file_content.decode(encoding)
            except (UnicodeDecodeError, AttributeError):
                # Fallback to UTF-8 with error replacement
                text = file_content.decode("utf-8", errors="replace")

            if not text.strip():
                raise ValueError("Empty file content")

            return text

        except Exception as e:
            logger.error(f"Error extracting text: {str(e)}")
            raise ValueError(f"Failed to extract text content: {str(e)}")

    def clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        # Remove excessive whitespace
        text = re.sub(r"\n{3,}", "\n\n", text)  # Max 2 consecutive newlines
        text = re.sub(r" {2,}", " ", text)  # Max 1 space

        # Remove PDF artifacts
        text = re.sub(r"\x0c", "\n", text)  # Form feed characters
        text = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]", "", text)  # Control chars

        return text.strip()

    def chunk_text(
        self,
        text: str,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
    ) -> list[dict]:
        """
        Chunk text using LangChain's RecursiveCharacterTextSplitter.

        Args:
            text: Text to chunk
            chunk_size: Maximum size of each chunk
            chunk_overlap: Overlap between chunks

        Returns:
            list[dict]: List of chunks with metadata
        """
        # Create splitter with custom parameters
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            is_separator_regex=False,
            separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""],
        )

        # Split text using LangChain
        text_chunks = splitter.split_text(text)

        # Convert to dict format with metadata
        chunks = []
        for index, chunk_text in enumerate(text_chunks):
            chunks.append(
                {
                    "content": chunk_text,
                    "index": index,
                    "char_count": len(chunk_text),
                }
            )

        logger.info(f"Split text into {len(chunks)} chunks using LangChain")
        return chunks

    def chunk_text_to_documents(
        self,
        text: str,
        metadata: dict | None = None,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
    ) -> list[Document]:
        """
        Chunk text into LangChain Document objects.

        Args:
            text: Text to chunk
            metadata: Optional metadata to attach to documents
            chunk_size: Maximum size of each chunk
            chunk_overlap: Overlap between chunks

        Returns:
            list[Document]: LangChain Document objects
        """
        # Create splitter
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            is_separator_regex=False,
            separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""],
        )

        # Create initial Document
        doc = Document(page_content=text, metadata=metadata or {})

        # Split into chunks (returns list of Documents)
        documents = splitter.split_documents([doc])

        logger.info(f"Created {len(documents)} LangChain Document chunks")
        return documents

    def process_file(
        self,
        file_content: bytes,
        filename: str,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
    ) -> dict:
        """
        Complete file processing pipeline using LangChain.

        Returns:
            dict with:
            - content: str - Full extracted text
            - chunks: list[dict] - Text chunks ready for embedding
            - documents: list[Document] - LangChain Document objects
            - metadata: dict - File metadata
        """
        # Validate
        is_valid, error_msg = self.validate_file_size(len(file_content))
        if not is_valid:
            raise ValueError(error_msg)

        if not self.is_supported(filename):
            raise ValueError(
                f"Unsupported file type. Supported: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )

        # Extract text
        raw_text = self.extract_text(file_content, filename)

        # Clean text
        cleaned_text = self.clean_text(raw_text)

        # Chunk text using LangChain
        chunks = self.chunk_text(cleaned_text, chunk_size, chunk_overlap)

        # Create LangChain Documents
        file_metadata = {
            "filename": filename,
            "file_size": len(file_content),
            "file_type": Path(filename).suffix.lower(),
        }
        documents = self.chunk_text_to_documents(
            cleaned_text, file_metadata, chunk_size, chunk_overlap
        )

        # Final metadata
        metadata = {
            **file_metadata,
            "total_chars": len(cleaned_text),
            "total_chunks": len(chunks),
        }

        logger.info(
            f"Processed {filename} with LangChain: {len(cleaned_text)} chars, {len(chunks)} chunks"
        )

        return {
            "content": cleaned_text,
            "chunks": chunks,
            "documents": documents,
            "metadata": metadata,
        }


# Global processor instance
document_processor = DocumentProcessor()
